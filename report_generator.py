from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re
from datetime import datetime
from fpdf import FPDF

def clean_markdown(text):
    """Удаляет markdown-разметку и нормализует эмодзи"""
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'^---+\s*$', '', text, flags=re.MULTILINE)
    text = text.replace('**', '')
    text = text.replace('*', '')
    text = text.replace('__', '')
    text = text.replace('_', '')
    text = text.replace('`', '')
    text = text.replace('\ufe0f', '')

    replacements = {
        '✌': '❌', '✖': '❌', '✕': '❌', '✘': '❌',
        '✗': '❌', '×': '❌', '❎': '❌',
        '✔': '✅', '✓': '✅',
        '☑': '✅',
        '☐': '❌',
    }
    for wrong, correct in replacements.items():
        text = text.replace(wrong, correct)

    return text.strip()


def create_contract_report(contract_data, output_dir):
    """Создаёт Word-отчёт"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading('АНАЛИЗ ДОГОВОРА', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contract_type = contract_data.get('contract_type', 'не определён')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'📋 Тип договора: {contract_type}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_heading('📄 Информация о документе', level=1)
    p = doc.add_paragraph()
    p.add_run('Файл: ').bold = True
    p.add_run(contract_data['filename'])

    doc.add_heading('📋 Реквизиты договора', level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    data_rows = [
        ('Номер договора', contract_data['number']),
        ('Дата заключения', contract_data['date']),
        ('Сумма договора', contract_data['amount']),
        ('Найдено ИНН', f"{len(contract_data['inn_list'])} шт. ({', '.join(contract_data['inn_list']) or 'не найдено'})"),
        ('Размер пеней', f"{contract_data['peni']} — {contract_data['peni_risk']}"),
    ]
    for i, (label, value) in enumerate(data_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[1].text = value

    doc.add_heading('⚠️ Оценка риска по пеням', level=1)
    if 'ВЫШЕ НОРМЫ' in contract_data['peni_risk']:
        p = doc.add_paragraph()
        run = p.add_run('🔴 ВЫСОКИЙ РИСК: ')
        run.bold = True
        run.font.color.rgb = RGBColor(156, 0, 6)
        p.add_run('Размер пеней превышает рыночную норму.')
    elif 'норма' in contract_data['peni_risk']:
        p = doc.add_paragraph()
        run = p.add_run('🟢 НОРМА: ')
        run.bold = True
        run.font.color.rgb = RGBColor(0, 97, 0)
        p.add_run('Размер пеней в пределах нормы.')
    else:
        doc.add_paragraph('⚪ Информация о пенях не найдена.')

    doc.add_heading('📋 Юридический чек-лист (ИИ-анализ)', level=1)
    ai_text = contract_data.get('ai_analysis', '')
    checklist_items = parse_checklist(ai_text)

    if checklist_items:
        checklist_table = doc.add_table(rows=len(checklist_items) + 1, cols=3)
        checklist_table.style = 'Light Grid Accent 1'
        headers = ['Статус', 'Пункт', 'Комментарий']
        for i, header in enumerate(headers):
            cell = checklist_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        for row_idx, item in enumerate(checklist_items, 1):
            status_cell = checklist_table.rows[row_idx].cells[0]
            status_cell.text = item['status']
            if '✅' in item['status']:
                status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
            elif '⚠' in item['status']:
                status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(196, 120, 0)
            elif '❌' in item['status']:
                status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(156, 0, 6)
            checklist_table.rows[row_idx].cells[1].text = item['title']
            checklist_table.rows[row_idx].cells[2].text = item['comment']

        doc.add_paragraph()
        summary = calculate_summary(checklist_items)
        if summary:
            p = doc.add_paragraph()
            p.add_run('📊 ИТОГ: ').bold = True
            p.add_run(summary)

    critical_risks = extract_critical_risks(ai_text)
    if critical_risks:
        doc.add_heading('🔥 Критические риски', level=1)
        for risk in critical_risks:
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(risk)
            run.font.color.rgb = RGBColor(156, 0, 6)

    doc.add_heading('✅ Чек-лист для исправления', level=1)
    checklist = generate_checklist(contract_data, checklist_items)
    for i, item in enumerate(checklist, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run('☐ ').font.size = Pt(14)
        p.add_run(item['task'])
        if item['priority'] == 'высокий':
            p.add_run(f"  [{item['priority'].upper()}]").bold = True
            p.runs[-1].font.color.rgb = RGBColor(156, 0, 6)
        else:
            p.add_run(f"  [{item['priority'].upper()}]")
            p.runs[-1].font.color.rgb = RGBColor(196, 120, 0)

    doc.add_paragraph()
    doc.add_paragraph('_' * 50)
    p = doc.add_paragraph()
    p.add_run('Отчёт сформирован автоматически системой ИИ-анализа договоров').italic = True
    p = doc.add_paragraph()
    p.add_run(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}').italic = True

    safe_filename = re.sub(r'[^\w\-.]', '_', contract_data['filename'].rsplit('.', 1)[0])
    output_path = Path(output_dir) / f'Отчёт_{safe_filename}.docx'
    doc.save(output_path)
    return output_path


def parse_checklist(ai_text):
    """ПРОСТОЙ И НАДЁЖНЫЙ парсер"""
    
    cleaned = clean_markdown(ai_text)
    cleaned = cleaned.replace('\ufe0f', '')
    
    standard_names = {
        1: 'Предмет договора',
        2: 'Цена и порядок расчётов',
        3: 'Сроки исполнения',
        4: 'Ответственность',
        5: 'Форс-мажор',
        6: 'Подсудность',
        7: 'Расторжение',
        8: 'Персональные данные',
        9: 'Существенные условия',
        10: 'Односторонние изменения'
    }
    
    items = []
    found_numbers = set()
    lines = cleaned.split('\n')
    
    for line in lines:
        line = line.strip()
        
        line_lower = line.lower()
        # Пропускаем служебные строки (continue, НЕ break!)
        if 'тип договора' in line_lower and ':' in line_lower:
            continue
        if 'субъектный состав' in line_lower:
            continue
        # Прерываем на рисках и итоге
        if 'критические риск' in line_lower:
            break
        if line_lower.startswith('итого') or ('итого' in line_lower and ':' in line_lower and 'из' in line_lower):
            break
        
        # Паттерн: НОМЕР. НАЗВАНИЕ: СТАТУС КОММЕНТАРИЙ
        match = re.match(r'^(\d{1,2})\.\s*(.+?):\s*([✅⚠❌])\s*(.+)$', line)
        
        if match:
            number = int(match.group(1))
            title_raw = match.group(2).strip()
            status = match.group(3)
            comment = match.group(4).strip()
            
            if 1 <= number <= 10 and number not in found_numbers:
                title = None
                for std_num, std_name in standard_names.items():
                    if std_name.lower() in title_raw.lower() or title_raw.lower() in std_name.lower():
                        title = std_name
                        break
                    if std_num == number:
                        keywords = [kw for kw in std_name.lower().split() if len(kw) > 3]
                        if any(kw in title_raw.lower() for kw in keywords):
                            title = std_name
                            break
                
                if not title:
                    title = standard_names.get(number, title_raw)
                
                comment_lower = comment.lower()
                
                absence_words = [
                    'отсутствует', 'отсутствуют', 'отсутствие', 'не указан', 'не указана',
                    'не указано', 'не упомянут', 'не определён', 'не содержится',
                    'не включён', 'не описан', 'не предусмотрен', 'полностью отсутствует',
                    'нет пункта', 'нет раздела', 'нет указания', 'не имеется',
                    'не прописан', 'не прописана', 'не прописано', 'не предусмотрен раздел'
                ]
                
                clarification_words = [
                    'требуется уточнение', 'требует уточнения', 'требует доработки',
                    'недостаточно подробно', 'не детализированы', 'рекомендуется уточнить',
                    'необходимо уточнить', 'частично указано', 'общие формулировки',
                    'не подробно', 'требует конкретизации', 'неясно', 'неполный',
                    'требует детализации', 'неполная информация', 'необходима детализация',
                    'не указан размер', 'не указаны сроки', 'нет конкретных'
                ]
                
                has_absence = any(word in comment_lower for word in absence_words)
                has_clarification = any(word in comment_lower for word in clarification_words)
                
                if has_absence and status != '❌':
                    status = '❌'
                elif has_clarification and status == '✅':
                    status = '⚠️'
                
                if comment and len(comment) > 3:
                    items.append({
                        'status': status,
                        'number': str(number),
                        'title': title,
                        'comment': comment
                    })
                    found_numbers.add(number)
    
    for num in range(1, 11):
        if num not in found_numbers:
            items.append({
                'status': '❌',
                'number': str(num),
                'title': standard_names[num],
                'comment': 'Пункт отсутствует в анализе ИИ. Требуется проверка вручную.'
            })
    
    items.sort(key=lambda x: int(x['number']))
    return items


def calculate_summary(checklist_items):
    """Автоматически пересчитывает итог"""
    if not checklist_items:
        return None
    ok_count = sum(1 for item in checklist_items if item['status'] == '✅')
    warn_count = sum(1 for item in checklist_items if item['status'] in ['⚠️', '⚠'])
    bad_count = sum(1 for item in checklist_items if item['status'] == '❌')
    total = len(checklist_items)
    return f"{ok_count} из {total} в порядке, {warn_count} требуют внимания, {bad_count} отсутствуют"


def extract_critical_risks(ai_text):
    """Парсер критических рисков"""
    risks = []
    lines = ai_text.split('\n')
    in_risks_block = False
    current_risk = ""

    for line in lines:
        line = line.strip()
        line = clean_markdown(line)

        if 'КРИТИЧЕСКИЕ РИСКИ' in line.upper():
            in_risks_block = True
            continue

        if in_risks_block:
            if line == '':
                if current_risk and len(current_risk) > 20:
                    risks.append(current_risk.strip())
                    current_risk = ""
                    if len(risks) >= 5:
                        break
                continue
            if re.match(r'^[✅⚠❌]\s*\d+\.', line):
                break
            if 'ИТОГО' in line.upper() or 'ТИП ДОГОВОРА' in line.upper():
                break

        if in_risks_block:
            risk_part = re.sub(r'^\d+[\.\)]\s*', '', line)
            risk_part = re.sub(r'^[✅⚠❌📍🔹🔸🔴🟢🟡➖]+\s*', '', risk_part)
            risk_part = risk_part.strip()

            junk_phrases = ['что именно опасно', 'нарушается статья', 'рекомендация:',
                           'описание риска', 'риск + статья', '[риск', '###']
            if any(junk in risk_part.lower() for junk in junk_phrases):
                continue

            if risk_part.startswith('—') or risk_part.startswith('-'):
                if current_risk:
                    risk_part = risk_part.lstrip('—-').strip()
                    current_risk += ' — ' + risk_part
            else:
                if current_risk and len(current_risk) > 20:
                    risks.append(current_risk.strip())
                current_risk = risk_part

            if len(risks) >= 5:
                break

    if current_risk and len(current_risk) > 20 and len(risks) < 5:
        risks.append(current_risk.strip())

    return risks[:5]


def generate_checklist(contract_data, checklist_items):
    """Генерирует чек-лист для исправления"""
    checklist = []

    if contract_data['number'] == 'не указан':
        checklist.append({'task': 'Указать номер договора в преамбуле', 'priority': 'высокий'})
    if contract_data['date'] == 'не указана':
        checklist.append({'task': 'Указать дату заключения договора', 'priority': 'высокий'})
    if contract_data['amount'] == 'не указана':
        price_item = next((i for i in checklist_items if 'цен' in i['title'].lower() or 'расч' in i['title'].lower()), None)
        if price_item and price_item['status'] in ['⚠️', '⚠', '❌']:
            checklist.append({'task': 'Чётко прописать сумму договора и порядок расчётов', 'priority': 'высокий'})
    if len(contract_data['inn_list']) < 2:
        checklist.append({'task': 'Проверить реквизиты сторон — должны быть ИНН обеих сторон', 'priority': 'высокий'})
    if 'ВЫШЕ НОРМЫ' in contract_data['peni_risk']:
        checklist.append({'task': 'Снизить размер пеней до рыночной нормы', 'priority': 'высокий'})

    for item in checklist_items:
        title_lower = item['title'].lower()
        status = item['status']
        if status == '✅':
            continue

        if 'форс' in title_lower:
            checklist.append({'task': 'Добавить/уточнить раздел о форс-мажорных обстоятельствах', 'priority': 'средний'})
        if 'подсудн' in title_lower or 'спор' in title_lower:
            checklist.append({'task': 'Добавить раздел о подсудности споров', 'priority': 'средний'})
        if 'расторж' in title_lower:
            checklist.append({'task': 'Уточнить условия расторжения договора', 'priority': 'средний'})
        if 'перс' in title_lower or '152' in title_lower:
            checklist.append({'task': 'Добавить согласие на обработку персональных данных (152-ФЗ)', 'priority': 'высокий'})
        if 'односторонн' in title_lower:
            checklist.append({'task': 'Уточнить условия односторонних изменений договора', 'priority': 'высокий'})
        if 'предмет' in title_lower:
            checklist.append({'task': 'Уточнить формулировки предмета договора', 'priority': 'высокий'})
        if 'срок' in title_lower:
            checklist.append({'task': 'Уточнить сроки исполнения обязательств', 'priority': 'средний'})
        if 'ответ' in title_lower:
            if not any('ответственност' in t['task'].lower() for t in checklist):
                checklist.append({'task': 'Добавить/уточнить раздел об ответственности сторон', 'priority': 'средний'})

    if not checklist:
        checklist = [
            {'task': 'Проверить соответствие реквизитов сторон', 'priority': 'средний'},
            {'task': 'Убедиться в наличии всех существенных условий', 'priority': 'средний'},
        ]

    return checklist


def generate_pdf_report(contract_data, output_dir):
    """Создаёт PDF-отчёт с анализом договора"""
    
    # === ИНИЦИАЛИЗАЦИЯ PDF ===
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # === ПОДКЛЮЧЕНИЕ ШРИФТА С КИРИЛЛИЦЕЙ ===
    # Путь к шрифтам Roboto
    font_regular = Path(__file__).parent / 'Roboto-Regular.ttf'
    font_bold = Path(__file__).parent / 'Roboto-Bold.ttf'
    
    if font_regular.exists() and font_bold.exists():
        pdf.add_font('Roboto', '', str(font_regular), uni=True)
        pdf.add_font('Roboto', 'B', str(font_bold), uni=True)
        pdf.set_font('Roboto', size=11)
    elif font_regular.exists():
        # Только обычный шрифт
        pdf.add_font('Roboto', '', str(font_regular), uni=True)
        pdf.set_font('Roboto', size=11)
    else:
        print(f"⚠️  Шрифты не найдены")
        pdf.set_font('Helvetica', size=11)
    
    # === ЗАГОЛОВОК ===
    pdf.set_font_size(18)
    pdf.set_text_color(68, 114, 196)
    pdf.cell(0, 15, 'АНАЛИЗ ДОГОВОРА', align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    
    # === ТИП ДОГОВОРА ===
    contract_type = contract_data.get('contract_type', 'не определён')
    pdf.set_font_size(14)
    pdf.set_text_color(68, 114, 196)
    pdf.cell(0, 10, f'Тип договора: {contract_type}', align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    
    # === ИНФОРМАЦИЯ О ДОКУМЕНТЕ ===
    pdf.set_text_color(0, 0, 0)
    pdf.set_font_size(14)
    pdf.cell(0, 10, 'Информация о документе', new_x="LMARGIN", new_y="NEXT")
    pdf.set_font_size(11)
    pdf.cell(0, 7, f'Файл: {contract_data["filename"]}', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    
    # === РЕКВИЗИТЫ ДОГОВОРА ===
    pdf.set_font_size(14)
    pdf.cell(0, 10, 'Реквизиты договора', new_x="LMARGIN", new_y="NEXT")
    pdf.set_font_size(11)
    
    # Таблица реквизитов
    requisites = [
        ('Номер договора', contract_data['number']),
        ('Дата заключения', contract_data['date']),
        ('Сумма договора', contract_data['amount']),
        ('Найдено ИНН', f"{len(contract_data['inn_list'])} шт. ({', '.join(contract_data['inn_list']) or 'не найдено'})"),
        ('Размер пеней', f"{contract_data['peni']} — {contract_data['peni_risk']}"),
    ]
    
    for label, value in requisites:
        pdf.set_font(style='B')
        pdf.cell(60, 7, label, border=1)
        pdf.set_font(style='')
        pdf.cell(0, 7, str(value), border=1, new_x="LMARGIN", new_y="NEXT")
    
    pdf.ln(3)
    
    # === ОЦЕНКА РИСКА ПО ПЕНЯМ ===
    pdf.set_font_size(14)
    pdf.cell(0, 10, 'Оценка риска по пеням', new_x="LMARGIN", new_y="NEXT")
    pdf.set_font_size(11)
    
    if 'ВЫШЕ НОРМЫ' in contract_data['peni_risk']:
        pdf.set_text_color(156, 0, 6)
        pdf.cell(0, 7, 'ВЫСОКИЙ РИСК: Размер пеней превышает рыночную норму.', new_x="LMARGIN", new_y="NEXT")
    elif 'норма' in contract_data['peni_risk']:
        pdf.set_text_color(0, 97, 0)
        pdf.cell(0, 7, 'НОРМА: Размер пеней в пределах нормы.', new_x="LMARGIN", new_y="NEXT")
    else:
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 7, 'Информация о пенях не найдена.', new_x="LMARGIN", new_y="NEXT")
    
    pdf.set_text_color(0, 0, 0)
    pdf.ln(3)
    
    # === ЮРИДИЧЕСКИЙ ЧЕК-ЛИСТ ===
    ai_text = contract_data.get('ai_analysis', '')
    checklist_items = parse_checklist(ai_text)
    
    pdf.set_font_size(14)
    pdf.cell(0, 10, 'Юридический чек-лист (ИИ-анализ)', new_x="LMARGIN", new_y="NEXT")
    
    if checklist_items:
        # Заголовок таблицы
        pdf.set_font_size(10)
        pdf.set_fill_color(68, 114, 196)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font(style='B')
        pdf.cell(15, 7, 'Статус', border=1, fill=True, align='C')
        pdf.cell(45, 7, 'Пункт', border=1, fill=True, align='C')
        pdf.cell(0, 7, 'Комментарий', border=1, fill=True, align='C', new_x="LMARGIN", new_y="NEXT")
        
        # Строки таблицы
        pdf.set_text_color(0, 0, 0)
        pdf.set_font(style='')
        
        for item in checklist_items:
            # Цвет статуса
            if '✅' in item['status']:
                pdf.set_text_color(0, 128, 0)
            elif '⚠' in item['status']:
                pdf.set_text_color(196, 120, 0)
            elif '❌' in item['status']:
                pdf.set_text_color(156, 0, 6)
            
            # Статус
            pdf.set_font_size(10)
            pdf.cell(15, 7, item['status'], border=1, align='C')
            
            # Пункт
            pdf.set_text_color(0, 0, 0)
            pdf.cell(45, 7, item['title'], border=1)
            
            # Комментарий (с переносом)
            comment = item['comment'][:150] + '...' if len(item['comment']) > 150 else item['comment']
            # Разбиваем длинный комментарий на строки
            pdf.multi_cell(0, 7, comment, border=1, new_x="LMARGIN", new_y="NEXT")
        
        # Итог
        pdf.ln(3)
        summary = calculate_summary(checklist_items)
        if summary:
            pdf.set_font_size(12)
            pdf.set_font(style='B')
            pdf.set_text_color(68, 114, 196)
            pdf.cell(0, 10, f'ИТОГ: {summary}', new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)
            pdf.set_font(style='')
    
    # === КРИТИЧЕСКИЕ РИСКИ ===
    critical_risks = extract_critical_risks(ai_text)
    if critical_risks:
        pdf.ln(3)
        pdf.set_font_size(14)
        pdf.set_text_color(156, 0, 6)
        pdf.cell(0, 10, 'Критические риски', new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.set_font_size(11)
        
        for i, risk in enumerate(critical_risks, 1):
            pdf.set_font(style='B')
            pdf.cell(8, 7, f'{i}.')
            pdf.set_font(style='')
            pdf.multi_cell(0, 7, risk, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
    
    # === ЧЕК-ЛИСТ ДЛЯ ИСПРАВЛЕНИЯ ===
    pdf.ln(3)
    pdf.set_font_size(14)
    pdf.set_text_color(0, 128, 0)
    pdf.cell(0, 10, 'Чек-лист для исправления', new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.set_font_size(11)
    
    checklist = generate_checklist(contract_data, checklist_items)
    for i, item in enumerate(checklist, 1):
        priority_mark = f"  [{item['priority'].upper()}]"
        pdf.cell(5, 7, f'{i}.')
        pdf.cell(5, 7, chr(9744))  # ☐ символ чекбокса
        pdf.multi_cell(0, 7, f"{item['task']}{priority_mark}", new_x="LMARGIN", new_y="NEXT")
    
    # === ПОДПИСЬ ===
    pdf.ln(5)
    pdf.set_font_size(9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 5, '_' * 80, new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 5, 'Отчёт сформирован автоматически системой ИИ-анализа договоров', new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 5, f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}', new_x="LMARGIN", new_y="NEXT")
    
    # === СОХРАНЕНИЕ ===
    safe_filename = re.sub(r'[^\w\-.]', '_', contract_data['filename'].rsplit('.', 1)[0])
    output_path = Path(output_dir) / f'Отчёт_{safe_filename}.pdf'
    pdf.output(str(output_path))
    
    return output_path
